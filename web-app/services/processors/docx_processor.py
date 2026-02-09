import time
from typing import List

from docx import Document as DocxDocument

from config.settings import LLM_MAX_INPUT_TOKENS
from utils.logger import get_logger

logger = get_logger(__name__)

# Approximate chars per token for estimation
CHARS_PER_TOKEN = 4


class DOCXProcessor:
    """
    Processes DOCX documents using python-docx library.
    Extracts text and tables, splits into batches respecting token limits.
    """

    def __init__(self):
        logger.info("[DOCXProcessor] Initialized (using python-docx)")

    def process(self, file_path: str) -> List[dict]:
        """
        Process DOCX file and extract text content.

        Args:
            file_path: Path to the DOCX file

        Returns:
            List of section dicts with 'page_number', 'content_type', 'text_content', etc.
        """
        logger.info(f"[DOCXProcessor] Processing DOCX...")
        start_time = time.time()

        doc = DocxDocument(file_path)

        # Extract all content (paragraphs and tables)
        all_content = self._extract_content(doc)

        # Split into token-safe sections
        sections = self._split_into_sections(all_content)

        elapsed = time.time() - start_time
        logger.info(f"[DOCXProcessor] Done in {elapsed:.2f}s - {len(sections)} sections")

        return sections

    def _extract_content(self, doc: DocxDocument) -> List[dict]:
        """
        Extract paragraphs and tables from document in order.

        Args:
            doc: python-docx Document object

        Returns:
            List of content items with type and text
        """
        content_items = []

        # Iterate through document body elements in order
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                # Find corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element is element:
                        text = para.text.strip()
                        if text:
                            # Check if it's a heading
                            style_name = para.style.name if para.style else ""
                            if "Heading" in style_name:
                                content_items.append({
                                    "type": "heading",
                                    "text": f"[{style_name.upper()}] {text}",
                                    "style": style_name
                                })
                            else:
                                content_items.append({
                                    "type": "paragraph",
                                    "text": text
                                })
                        break

            elif element.tag.endswith('tbl'):  # Table
                for table in doc.tables:
                    if table._element is element:
                        table_text = self._format_table(table)
                        if table_text:
                            content_items.append({
                                "type": "table",
                                "text": table_text
                            })
                        break

        logger.info(f"[DOCXProcessor] Extracted {len(content_items)} content items")
        return content_items

    def _format_table(self, table) -> str:
        """Format a table into readable text."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Skip empty rows
            if any(cells):
                rows.append(" | ".join(cells))

        if rows:
            return "TABLE:\n" + "\n".join(rows)
        return ""

    def _split_into_sections(self, content_items: List[dict]) -> List[dict]:
        """
        Split content into token-safe sections.

        Args:
            content_items: List of content items from extraction

        Returns:
            List of section dicts ready for LLM processing
        """
        sections = []
        current_content = []
        current_tokens = 0
        section_num = 1

        # Calculate safe token limit (leave room for system prompt)
        safe_token_limit = LLM_MAX_INPUT_TOKENS - 2000

        def flush_section():
            nonlocal current_content, current_tokens, section_num
            if current_content:
                text = "\n\n".join(current_content)
                sections.append({
                    "page_number": section_num,
                    "content_type": "text",
                    "text_content": text,
                    "char_count": len(text),
                    "estimated_tokens": len(text) // CHARS_PER_TOKEN
                })
                logger.info(f"  Section {section_num}: {len(text)} chars, ~{len(text) // CHARS_PER_TOKEN} tokens")
                section_num += 1
                current_content = []
                current_tokens = 0

        for item in content_items:
            text = item["text"]
            item_tokens = len(text) // CHARS_PER_TOKEN + 1

            # Check if adding this item would exceed limit
            if current_tokens + item_tokens > safe_token_limit:
                flush_section()

            current_content.append(text)
            current_tokens += item_tokens

        # Flush remaining content
        flush_section()

        # If no content extracted
        if not sections:
            sections.append({
                "page_number": 1,
                "content_type": "text",
                "text_content": "[Empty document]",
                "char_count": 16,
                "estimated_tokens": 4
            })

        return sections
